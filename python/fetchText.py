#!/usr/bin/python
#coding=utf-8

import urllib.request;
import os;
from os import system;
from lxml import etree;
import ssl;
import re;
import time;
from docx import Document;

ssl._create_default_https_context = ssl._create_unverified_context

header = {"User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/52.0.2743.116 Safari/537.36", "Connection": "keep-alive"}

fetchUrl = 'http://www.fzl258.com/index.php/arttype/qingganxiaoshuo';
host = 'http://www.fzl258.com';

def getTotalNum ():
    if not os.path.exists('textdir'):
        os.makedirs('textdir')
    
    if os.path.exists('textlog.txt'):
        os.remove('textlog.txt');

    for i in range(1, 91):
        if (i == 1):
            url = fetchUrl + '.html'
        else:
            url = fetchUrl + '-' + str(i) +  '.html'

        getRankNum(url, i)

def getRankNum (url, index):
    # maxTryNum = 20
    # for tries in range(maxTryNum):
    #     try:
            
    #     except:
    #         if tries < (maxTryNum - 1):
    #             continue
    #         else:
    #             print("Has tried %d times to access url %s, all failed!" % (maxTryNum, url))
    #             break

    req = urllib.request.Request(url, headers=header);
    html = urllib.request.urlopen(req);
    htmldata = html.read();
    htmlPath = etree.HTML(htmldata);

    pages = htmlPath.xpath('//div[@class="box list channel"]/ul/li/a/@href');
    print(url)
    # print(pages)

    for i in range(len(pages)):
        time.sleep(3)

        try:
            getItem(host + pages[i], index, i)
        except Exception:
            print("Has tried  all failed!")
            pass;

def getItem(url, index, tindex):
    req = urllib.request.Request(url, headers=header);
    html = urllib.request.urlopen(req);
    htmldata = html.read();
    htmlPath = etree.HTML(htmldata);
    # print(url)
    # 

    title = htmlPath.xpath('//div[@class="page_title"]/text()');
    text = htmlPath.xpath('//div[@class="w685"]/text() | //div[@class="w685"]/p/text()');
    
    # print(title)
    # print(text)

    file = Document();
    file.add_heading(title[0])
    for i in range(0, len(text)):
        file.add_paragraph(text[i]);
    file.save('textdir/%d_%d_%s.docx' % (index, tindex, title[0]))

    file_handle = open('textlog.txt', mode='a')
    file_handle.write('%d_%d_%s   %d' % (index, tindex, title[0], len(text)) + '\n')
    
    file_handle.close()

if __name__ == '__main__':
    getTotalNum();

    print('finish');
    system('shutdown -s');